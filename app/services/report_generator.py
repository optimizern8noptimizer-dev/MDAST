import json
import logging
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

from app.config import Config
from app.database import SessionLocal
from app.models.scan import Scan
from app.models.finding import Finding, Severity
from app.models.report import Report, ReportFormat
from app.models.user import User

logger = logging.getLogger(__name__)

COLOR = {
    "critical": ((0.78, 0.08, 0.08), "C71515"),
    "high":     ((0.85, 0.33, 0.0),  "D95400"),
    "medium":   ((0.85, 0.65, 0.0),  "D9A600"),
    "low":      ((0.12, 0.53, 0.12), "1F871F"),
    "info":     ((0.2,  0.4,  0.7),  "336699"),
    "header":   ((0.11, 0.27, 0.53), "1C4587"),
    "white":    ((1.0,  1.0,  1.0),  "FFFFFF"),
    "light":    ((0.95, 0.95, 0.97), "F2F2F8"),
}


def generate_report(scan_id: int, generated_by_user_id: int, fmt: str = "pdf") -> Report:
    db = SessionLocal()
    try:
        scan = db.query(Scan).filter(Scan.id == scan_id).first()
        if not scan:
            raise ValueError(f"Scan {scan_id} not found")

        findings = (db.query(Finding)
                    .filter(Finding.scan_id == scan_id)
                    .order_by(Finding.cvss_score.desc())
                    .all())

        user      = db.query(User).filter(User.id == scan.user_id).first()
        gen_user  = db.query(User).filter(User.id == generated_by_user_id).first()

        report_data = _build_report_data(scan, findings, user, gen_user)

        reports_dir = Path(Config.REPORTS_FOLDER)
        reports_dir.mkdir(parents=True, exist_ok=True)
        ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        filename  = f"MASP_Report_Scan{scan_id}_{ts}.{fmt}"
        file_path = reports_dir / filename

        if fmt == "pdf":
            _generate_pdf(report_data, str(file_path))
        elif fmt == "docx":
            _generate_docx(report_data, str(file_path))
        else:
            raise ValueError(f"Unsupported format: {fmt}")

        file_size = file_path.stat().st_size

        report = Report(
            scan_id=scan_id,
            generated_by=generated_by_user_id,
            format=ReportFormat(fmt),
            file_path=str(file_path),
            file_size_bytes=file_size,
        )
        db.add(report)
        db.commit()
        db.refresh(report)
        return report
    finally:
        db.close()


def _build_report_data(scan, findings, analyst, generator) -> dict:
    sev_counts = Counter(f.severity.value for f in findings)
    risk_level = (
        "CRITICAL" if sev_counts.get("critical", 0) > 0 else
        "HIGH"     if sev_counts.get("high", 0) > 0     else
        "MEDIUM"   if sev_counts.get("medium", 0) > 0   else
        "LOW"
    )
    return {
        "generated_at":  datetime.now(timezone.utc).strftime("%d.%m.%Y %H:%M UTC"),
        "scan_id":       scan.id,
        "apk_name":      scan.apk_name,
        "package_name":  scan.package_name or "N/A",
        "scan_type":     scan.scan_type.value.upper(),
        "scan_status":   scan.status.value,
        "duration_sec":  scan.duration_seconds or 0,
        "analyst":       analyst.username if analyst else "N/A",
        "generator":     generator.username if generator else "N/A",
        "findings":      findings,
        "total_findings": len(findings),
        "sev_counts":    sev_counts,
        "sast_count":    sum(1 for f in findings if f.scan_source == "sast"),
        "dast_count":    sum(1 for f in findings if f.scan_source == "dast"),
        "top5":          sorted(findings, key=lambda f: f.cvss_score or 0, reverse=True)[:5],
        "risk_level":    risk_level,
        "critical_list": [f for f in findings if f.severity.value == "critical"],
        "high_list":     [f for f in findings if f.severity.value == "high"],
        "medium_list":   [f for f in findings if f.severity.value == "medium"],
        "standards": [
            "OWASP Mobile Top 10 2024",
            "OWASP MASVS v2.0",
            "CWE/SANS Top 25",
            "CVSS v3.1",
            "PCI DSS 4.0",
        ],
    }


def _generate_pdf(data: dict, output_path: str):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table,
        TableStyle, PageBreak, HRFlowable,
    )
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2.5*cm, bottomMargin=2*cm,
    )

    HDR  = colors.HexColor("#1C4587")
    SEV  = {
        "critical": colors.HexColor("#C71515"),
        "high":     colors.HexColor("#D95400"),
        "medium":   colors.HexColor("#D9A600"),
        "low":      colors.HexColor("#1F871F"),
        "info":     colors.HexColor("#336699"),
    }
    S = {
        "title":  ParagraphStyle("t",  fontSize=22, textColor=HDR, alignment=TA_CENTER, fontName="Helvetica-Bold", spaceAfter=6),
        "h1":     ParagraphStyle("h1", fontSize=14, textColor=HDR, spaceBefore=14, spaceAfter=6, fontName="Helvetica-Bold"),
        "h2":     ParagraphStyle("h2", fontSize=11, textColor=HDR, spaceBefore=10, spaceAfter=4, fontName="Helvetica-Bold"),
        "body":   ParagraphStyle("b",  fontSize=9,  spaceAfter=4, leading=14, alignment=TA_JUSTIFY),
        "code":   ParagraphStyle("c",  fontSize=7.5, fontName="Courier", backColor=colors.HexColor("#F2F2F8"), leftIndent=8, spaceAfter=5, leading=11),
    }

    story = []

    # Титульная
    story.append(Spacer(1, 2*cm))
    story.append(Paragraph("MOBILE APPLICATION SECURITY REPORT", S["title"]))
    story.append(Spacer(1, 0.4*cm))
    story.append(HRFlowable(width="100%", thickness=2, color=HDR))
    story.append(Spacer(1, 0.4*cm))

    meta = [
        ["Application", data["apk_name"]],
        ["Package",     data["package_name"]],
        ["Scan Type",   data["scan_type"]],
        ["Scan ID",     f"#{data['scan_id']}"],
        ["Analyst",     data["analyst"]],
        ["Generated",   data["generated_at"]],
        ["Overall Risk",data["risk_level"]],
    ]
    t = Table(meta, colWidths=[4.5*cm, 12*cm])
    t.setStyle(TableStyle([
        ("FONTNAME",  (0,0), (0,-1), "Helvetica-Bold"),
        ("FONTSIZE",  (0,0), (-1,-1), 9),
        ("TEXTCOLOR", (0,0), (0,-1), HDR),
        ("GRID",      (0,0), (-1,-1), 0.4, colors.HexColor("#CCCCCC")),
        ("ROWBACKGROUNDS", (0,0), (-1,-1), [colors.white, colors.HexColor("#F9F9FB")]),
        ("TEXTCOLOR", (1,6), (1,6), SEV.get(data["risk_level"].lower(), colors.grey)),
        ("FONTNAME",  (1,6), (1,6), "Helvetica-Bold"),
    ]))
    story.append(t)
    story.append(PageBreak())

    # Executive Summary
    story.append(Paragraph("1. Executive Summary", S["h1"]))
    story.append(HRFlowable(width="100%", thickness=1, color=HDR))
    sc = data["sev_counts"]
    story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph(
        f"Security analysis of <b>{data['apk_name']}</b> identified "
        f"<b>{data['total_findings']}</b> findings: "
        f"<font color='#C71515'><b>{sc.get('critical',0)} Critical</b></font>, "
        f"<font color='#D95400'><b>{sc.get('high',0)} High</b></font>, "
        f"<font color='#D9A600'>{sc.get('medium',0)} Medium</font>, "
        f"{sc.get('low',0)} Low. Overall risk: <b>{data['risk_level']}</b>.",
        S["body"]
    ))

    sum_rows = [
        ["Severity", "Count"],
        ["Critical", str(sc.get("critical",0))],
        ["High",     str(sc.get("high",0))],
        ["Medium",   str(sc.get("medium",0))],
        ["Low",      str(sc.get("low",0))],
        ["SAST",     str(data["sast_count"])],
        ["DAST",     str(data["dast_count"])],
    ]
    t2 = Table(sum_rows, colWidths=[6*cm, 3*cm])
    t2.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), HDR),
        ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
        ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0), (-1,-1), 9),
        ("GRID",       (0,0), (-1,-1), 0.4, colors.HexColor("#CCCCCC")),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#F2F2F8")]),
        ("TEXTCOLOR", (1,1),(1,1), SEV["critical"]),
        ("TEXTCOLOR", (1,2),(1,2), SEV["high"]),
        ("TEXTCOLOR", (1,3),(1,3), SEV["medium"]),
        ("FONTNAME",  (1,1),(1,3), "Helvetica-Bold"),
    ]))
    story.append(Spacer(1, 0.2*cm))
    story.append(t2)
    story.append(PageBreak())

    # Методология
    story.append(Paragraph("2. Methodology & Standards", S["h1"]))
    story.append(HRFlowable(width="100%", thickness=1, color=HDR))
    for std in data["standards"]:
        story.append(Paragraph(f"• {std}", S["body"]))

    # Детальные находки
    story.append(PageBreak())
    story.append(Paragraph("3. Detailed Findings", S["h1"]))
    story.append(HRFlowable(width="100%", thickness=1, color=HDR))

    if not data["findings"]:
        story.append(Paragraph("No security findings identified.", S["body"]))
    else:
        for i, f in enumerate(data["findings"], 1):
            sev = f.severity.value
            story.append(Spacer(1, 0.3*cm))
            story.append(Paragraph(
                f"<font color='#{COLOR[sev][1]}'>[{sev.upper()}]</font> "
                f"<b>#{i}: {f.title}</b>", S["h2"]
            ))
            mrows = [
                ["CWE", f.cwe_id or "N/A", "CVSS", str(f.cvss_score or 0)],
                ["OWASP", f.owasp_mobile or "N/A", "MASVS", f.masvs_id or "N/A"],
                ["PCI DSS", f.pci_dss_req or "N/A", "Source", f.scan_source.upper()],
            ]
            tm = Table(mrows, colWidths=[2.5*cm, 6*cm, 2.5*cm, 5.5*cm])
            tm.setStyle(TableStyle([
                ("FONTNAME", (0,0),(0,-1), "Helvetica-Bold"),
                ("FONTNAME", (2,0),(2,-1), "Helvetica-Bold"),
                ("FONTSIZE", (0,0),(-1,-1), 8),
                ("TEXTCOLOR",(0,0),(0,-1), HDR),
                ("TEXTCOLOR",(2,0),(2,-1), HDR),
                ("GRID",     (0,0),(-1,-1), 0.3, colors.HexColor("#DDDDDD")),
                ("BACKGROUND",(0,0),(-1,-1), colors.HexColor("#F9F9FB")),
            ]))
            story.append(tm)
            story.append(Spacer(1, 0.15*cm))
            story.append(Paragraph(f"<b>Description:</b> {f.description}", S["body"]))
            if f.file_path:
                story.append(Paragraph(
                    f"<b>Location:</b> {f.file_path}" + (f":{f.line_number}" if f.line_number else ""),
                    S["body"]
                ))
            if f.code_snippet:
                story.append(Paragraph(f.code_snippet[:300], S["code"]))
            story.append(Paragraph(f"<b>Recommendation:</b> {f.recommendation}", S["body"]))

    # Рекомендации
    story.append(PageBreak())
    story.append(Paragraph("4. Prioritized Recommendations", S["h1"]))
    story.append(HRFlowable(width="100%", thickness=1, color=HDR))
    for label, items, clr in [
        ("Immediate (24h)",    data["critical_list"], "critical"),
        ("Short-term (7d)",    data["high_list"],     "high"),
        ("Medium-term (30d)",  data["medium_list"],   "medium"),
    ]:
        if items:
            story.append(Paragraph(
                f"<font color='#{COLOR[clr][1]}'><b>{label}:</b></font>", S["h2"]
            ))
            for f in items[:10]:
                story.append(Paragraph(f"• [{f.cwe_id}] {f.title}: {f.recommendation}", S["body"]))

    # Заключение
    story.append(PageBreak())
    story.append(Paragraph("5. Conclusion", S["h1"]))
    story.append(HRFlowable(width="100%", thickness=1, color=HDR))
    story.append(Paragraph(
        f"Assessment of <b>{data['apk_name']}</b> revealed {data['total_findings']} findings. "
        f"Immediate remediation required for {sc.get('critical',0)} critical and "
        f"{sc.get('high',0)} high severity issues before production deployment.",
        S["body"]
    ))

    # Приложение
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph("Appendix: Regulatory References", S["h1"]))
    story.append(HRFlowable(width="100%", thickness=1, color=HDR))
    refs = [
        ["Standard", "URL"],
        ["OWASP MASVS v2.0",        "https://mas.owasp.org/MASVS/"],
        ["OWASP Mobile Top 10 2024","https://owasp.org/www-project-mobile-top-10/"],
        ["CWE/SANS Top 25",         "https://cwe.mitre.org/top25/"],
        ["PCI DSS 4.0",             "https://www.pcisecuritystandards.org/"],
        ["CVSS v3.1",               "https://www.first.org/cvss/"],
    ]
    tr = Table(refs, colWidths=[6*cm, 10*cm])
    tr.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0), HDR),
        ("TEXTCOLOR", (0,0),(-1,0), colors.white),
        ("FONTNAME",  (0,0),(-1,0), "Helvetica-Bold"),
        ("FONTSIZE",  (0,0),(-1,-1), 8),
        ("GRID",      (0,0),(-1,-1), 0.4, colors.HexColor("#CCCCCC")),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#F2F2F8")]),
    ]))
    story.append(Spacer(1, 0.2*cm))
    story.append(tr)

    doc.build(story)


def _generate_docx(data: dict, output_path: str):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def hex_color(h):
        return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

    def set_cell_bg(cell, h):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), h)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    HDR = hex_color(COLOR["header"][1])

    def add_h(text, level=1):
        p = doc.add_heading(text, level=level)
        for r in p.runs:
            r.font.color.rgb = HDR
        return p

    # Титульная
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOBILE APPLICATION SECURITY REPORT")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = HDR

    doc.add_paragraph()
    meta = [
        ("Application", data["apk_name"]),
        ("Package",     data["package_name"]),
        ("Scan Type",   data["scan_type"]),
        ("Scan ID",     f"#{data['scan_id']}"),
        ("Analyst",     data["analyst"]),
        ("Generated",   data["generated_at"]),
        ("Overall Risk",data["risk_level"]),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(meta):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = HDR
        set_cell_bg(t.rows[i].cells[0], "EAF0FB")

    doc.add_page_break()

    # Executive Summary
    add_h("1. Executive Summary")
    sc = data["sev_counts"]
    doc.add_paragraph(
        f"Analysis of '{data['apk_name']}' found {data['total_findings']} issues: "
        f"{sc.get('critical',0)} Critical, {sc.get('high',0)} High, "
        f"{sc.get('medium',0)} Medium, {sc.get('low',0)} Low. "
        f"Risk level: {data['risk_level']}."
    )

    # Методология
    doc.add_page_break()
    add_h("2. Methodology & Standards")
    for std in data["standards"]:
        doc.add_paragraph(std, style="List Bullet")

    # Находки
    doc.add_page_break()
    add_h("3. Detailed Findings")
    if not data["findings"]:
        doc.add_paragraph("No findings identified.")
    else:
        for i, f in enumerate(data["findings"], 1):
            add_h(f"Finding #{i}: {f.title}", 2)
            p = doc.add_paragraph()
            r = p.add_run(f"[{f.severity.value.upper()}]  ")
            r.bold = True
            r.font.color.rgb = hex_color(COLOR.get(f.severity.value, COLOR["info"])[1])

            mf = doc.add_table(rows=3, cols=4)
            mf.style = "Table Grid"
            items = [
                ("CWE", f.cwe_id or "N/A"), ("CVSS", str(f.cvss_score or 0)),
                ("OWASP", f.owasp_mobile or "N/A"), ("MASVS", f.masvs_id or "N/A"),
                ("PCI DSS", f.pci_dss_req or "N/A"), ("Source", f.scan_source.upper()),
            ]
            for idx, (k, v) in enumerate(items):
                row = mf.rows[idx // 2]
                ck, cv = (idx % 2)*2, (idx % 2)*2+1
                row.cells[ck].text = k
                row.cells[cv].text = v
                row.cells[ck].paragraphs[0].runs[0].bold = True
                row.cells[ck].paragraphs[0].runs[0].font.color.rgb = HDR
                set_cell_bg(row.cells[ck], "F0F4FB")

            doc.add_paragraph()
            doc.add_paragraph(f"Description: {f.description}")
            if f.file_path:
                doc.add_paragraph(f"Location: {f.file_path}" + (f":{f.line_number}" if f.line_number else ""))
            if f.code_snippet:
                p = doc.add_paragraph()
                r = p.add_run(f.code_snippet[:300])
                r.font.name = "Courier New"; r.font.size = Pt(8)
            doc.add_paragraph(f"Recommendation: {f.recommendation}")
            doc.add_paragraph()

    # Рекомендации
    doc.add_page_break()
    add_h("4. Prioritized Recommendations")
    for label, items, clr in [
        ("Immediate (24h)",   data["critical_list"], "critical"),
        ("Short-term (7d)",   data["high_list"],     "high"),
        ("Medium-term (30d)", data["medium_list"],   "medium"),
    ]:
        if items:
            p = doc.add_paragraph()
            r = p.add_run(label)
            r.bold = True; r.font.color.rgb = hex_color(COLOR[clr][1])
            for f in items[:10]:
                doc.add_paragraph(f"[{f.cwe_id}] {f.title}: {f.recommendation}", style="List Bullet")

    # Заключение
    doc.add_page_break()
    add_h("5. Conclusion")
    doc.add_paragraph(
        f"Assessment of '{data['apk_name']}' revealed {data['total_findings']} findings. "
        f"Immediate action required for {sc.get('critical',0)} critical and "
        f"{sc.get('high',0)} high severity issues."
    )

    doc.save(output_path)